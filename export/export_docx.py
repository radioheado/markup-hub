from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from pathlib import Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export numbered Markdown to a styled DOCX document."
    )
    parser.add_argument(
        "--input",
        default="build/merged/paper_numbered.md",
        help="Path to the numbered Markdown input file. Defaults to build/merged/paper_numbered.md.",
    )
    parser.add_argument(
        "--output",
        default="build/docx/paper.docx",
        help="Path to the DOCX output file. Defaults to build/docx/paper.docx.",
    )
    parser.add_argument(
        "--template",
        default="src/archive/3.24.2026_CNDCG_and_WCI.docx",
        help="Reference DOCX whose styles, colors, and page setup should be reused.",
    )
    parser.add_argument(
        "--audit-bold-leads",
        action="store_true",
        help="Accepted for compatibility with the prior exporter; ignored by the Pandoc-based path.",
    )
    return parser.parse_args()


def split_metadata_comment(lines: list[str]) -> tuple[list[str], list[str]]:
    if not lines or lines[0].strip() != "<!--":
        return [], lines
    comment: list[str] = []
    remaining = lines[:]
    while remaining:
        line = remaining.pop(0)
        comment.append(line)
        if line.strip() == "-->":
            break
    return comment, remaining


def extract_yaml_front_matter(comment_lines: list[str]) -> str:
    if not comment_lines:
        return ""
    yaml_lines: list[str] = []
    in_yaml = False
    for line in comment_lines:
        stripped = line.strip()
        if stripped == "---":
            yaml_lines.append("---")
            in_yaml = not in_yaml
            continue
        if in_yaml:
            yaml_lines.append(line)
    if yaml_lines.count("---") < 2:
        return ""
    return "\n".join(yaml_lines).strip() + "\n"


def normalize_heading_spacing(lines: list[str]) -> list[str]:
    normalized: list[str] = []
    for line in lines:
        stripped = line.strip()
        is_heading = bool(re.match(r"^#{1,6}\s+", stripped))
        if is_heading and normalized and normalized[-1] != "":
            normalized.append("")
        normalized.append(line)
    return normalized


def convert_inline_code_spans(lines: list[str]) -> list[str]:
    converted: list[str] = []
    for line in lines:
        converted.append(
            re.sub(
                r"`([^`\n]+)`",
                r'[\1]{custom-style="Verbatim Char"}',
                line,
            )
        )
    return converted


def build_pandoc_inputs(source_text: str) -> tuple[str, str]:
    lines = source_text.splitlines()
    metadata_comment, body_lines = split_metadata_comment(lines)
    yaml_front_matter = extract_yaml_front_matter(metadata_comment)
    normalized_body_lines: list[str] = []
    for index, line in enumerate(body_lines):
        normalized_body_lines.append(line)
        next_line = body_lines[index + 1] if index + 1 < len(body_lines) else ""
        if (
            re.match(r"^\*\*(Table \d+\.|Figure \d+\.)\*\*", line.strip())
            and next_line.strip().startswith("|")
        ):
            normalized_body_lines.append("")
    normalized_body_lines = normalize_heading_spacing(normalized_body_lines)
    normalized_body_lines = convert_inline_code_spans(normalized_body_lines)
    body = "\n".join(normalized_body_lines).lstrip("\n")
    return yaml_front_matter, body.rstrip() + "\n"


Document = None
WD_TABLE_ALIGNMENT = None
WD_CELL_VERTICAL_ALIGNMENT = None
WD_PARAGRAPH_ALIGNMENT = None
OxmlElement = None
parse_xml = None
nsdecls = None
qn = None
Pt = None
RGBColor = None
Mm = None
Inches = None


def set_run_font(run, name: str, size: int, *, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before_pt: float, after_pt: float = 0) -> None:
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)


def shade_run(run, fill: str) -> None:
    run_pr = run._r.get_or_add_rPr()
    existing = run_pr.find(qn("w:shd"))
    if existing is not None:
        run_pr.remove(existing)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    run_pr.append(shading)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    tc_pr.append(shading)


def set_cell_borders(cell, color: str = "2E75B6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "2")
        element.set(qn("w:color"), color)


def set_header_cell_borders(cell, *, is_first: bool, is_last: bool, color: str = "2E75B6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    edge_settings = {
        "top": ("single", color),
        "bottom": ("single", color),
        "left": ("single", color) if is_first else ("nil", color),
        "right": ("single", color) if is_last else ("nil", color),
    }
    for edge, (value, edge_color) in edge_settings.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        if value == "nil":
            for attr in ("w:sz", "w:color"):
                qname_attr = qn(attr)
                if qname_attr in element.attrib:
                    del element.attrib[qname_attr]
            continue
        element.set(qn("w:sz"), "2")
        element.set(qn("w:color"), edge_color)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", "80"), ("bottom", "80"), ("left", "120"), ("right", "120")):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), value)
        element.set(qn("w:type"), "dxa")
        tc_mar.append(element)
    tc_pr.append(tc_mar)


def get_table_column_widths(col_count: int, headers: list[str] | None = None) -> list[int]:
    if col_count <= 1:
        return [9026]
    normalized_headers = [(header or "").strip().lower() for header in (headers or [])]
    if len(normalized_headers) == col_count:
        weights: list[float] = []
        for header in normalized_headers:
            if header in {"item", "candidate"}:
                weights.append(3.6)
            elif header.startswith("case "):
                weights.append(2.8)
            elif header in {"calculation", "contribution", "question", "formula", "value"}:
                weights.append(2.2)
            elif header in {"rank", "i", "rel", "gold?", "record?", "p@k", "r@k", "rr score"}:
                weights.append(1.2)
            elif header in {"d(i)", "precision@k", "in gold set?"}:
                weights.append(1.6)
            else:
                weights.append(2.0)

        total_weight = sum(weights)
        widths = [max(900, int(9026 * weight / total_weight)) for weight in weights]
        width_delta = 9026 - sum(widths)
        if width_delta != 0:
            target_index = max(range(len(widths)), key=lambda index: weights[index])
            widths[target_index] += width_delta
        return widths

    if col_count == 2:
        return [2600, 6426]
    if col_count == 3:
        return [2200, 3413, 3413]
    base = 9026 // col_count
    widths = [base] * col_count
    widths[-1] += 9026 - sum(widths)
    return widths


def configure_table_width(table, col_widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9026")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    insert_at = 1 if tbl.tblPr is not None else 0
    tbl.insert(insert_at, tbl_grid)


def style_heading(paragraph, level: int) -> None:
    style_map = {
        1: (16, "2E74B5", 18),
        2: (13, "2E74B5", 12),
        3: (11, "4472C4", 8),
    }
    size, color, before = style_map[level]
    set_spacing(paragraph, before, 0)
    for run in paragraph.runs:
        set_run_font(run, "Calibri", size, bold=True, italic=False, color=color)


def style_caption(paragraph) -> None:
    set_spacing(paragraph, 12, 0)
    match = re.match(r"^(Table \d+\.|Figure \d+\.)(.*)$", paragraph.text.strip())
    if not match:
        for run in paragraph.runs:
            set_run_font(run, "Calibri", 9, italic=False)
        return
    label, rest = match.groups()
    paragraph.clear()
    first = paragraph.add_run(label)
    set_run_font(first, "Calibri", 9, bold=True, italic=False)
    if rest:
        second = paragraph.add_run(rest)
        set_run_font(second, "Calibri", 9, bold=False, italic=True)


def style_inline_code_run(run, size: int = 10) -> None:
    set_run_font(run, "Consolas", size, bold=False, italic=False, color="C7254E")
    shade_run(run, "F7F7F9")


def style_normal_paragraph(paragraph) -> None:
    set_spacing(paragraph, 6, 0)
    for run in paragraph.runs:
        if run.font.name == "Cambria Math":
            continue
        run_style_name = run.style.name if run.style else ""
        if run_style_name == "Verbatim Char":
            style_inline_code_run(run, 10)
            continue
        if run.text in {"✓", "✗"}:
            color = "3A7C22" if run.text == "✓" else "C00000"
            set_run_font(run, "Segoe UI Symbol", 11, bold=True, italic=False, color=color)
            continue
        set_run_font(run, "Calibri", 11, bold=bool(run.bold), italic=bool(run.italic))


def style_table_cell(cell, *, bold: bool = False, color: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        set_spacing(paragraph, 3, 3)
        for run in paragraph.runs:
            if run.font.name == "Cambria Math":
                continue
            run_style_name = run.style.name if run.style else ""
            if run_style_name == "Verbatim Char":
                style_inline_code_run(run, 9)
                continue
            if run.text in {"✓", "✗"}:
                symbol_color = "3A7C22" if run.text == "✓" else "C00000"
                set_run_font(run, "Segoe UI Symbol", 10, bold=True, italic=False, color=symbol_color)
                continue
            set_run_font(
                run,
                "Calibri",
                10,
                bold=bold or bool(run.bold),
                italic=bool(run.italic),
                color=color,
            )


def postprocess_docx(output_path: Path) -> None:
    document = Document(output_path)
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(0)

    title_style = document.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(16)
    title_style.font.bold = True

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name
        text = paragraph.text.strip()
        if style_name == "Heading 1":
            style_heading(paragraph, 1)
        elif style_name == "Heading 2":
            style_heading(paragraph, 2)
        elif style_name == "Heading 3":
            style_heading(paragraph, 3)
        elif style_name == "Title":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_spacing(paragraph, 0, 0)
        elif text.startswith(("Table ", "Figure ")):
            style_caption(paragraph)
        else:
            style_normal_paragraph(paragraph)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        try:
            table.style = "Normal Table"
        except KeyError:
            table.style = "Table Grid"
        header_cells = table.rows[0].cells if table.rows else []
        header_texts = [cell.text for cell in header_cells]
        col_widths = get_table_column_widths(len(table.columns), header_texts)
        configure_table_width(table, col_widths)
        for col_index, cell in enumerate(table.rows[0].cells):
            set_cell_width(cell, col_widths[col_index])
            set_cell_margins(cell)
            shade_cell(cell, "2E75B6")
            set_header_cell_borders(
                cell,
                is_first=col_index == 0,
                is_last=col_index == len(table.columns) - 1,
            )
            style_table_cell(cell, bold=True, color="FFFFFF")
        for row in table.rows[1:]:
            for col_index, cell in enumerate(row.cells):
                set_cell_width(cell, col_widths[col_index])
                set_cell_margins(cell)
                shade_cell(cell, "FFFFFF")
                set_cell_borders(cell)
                style_table_cell(cell)

    document.save(output_path)


def main() -> int:
    args = parse_args()
    global Document, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
    global OxmlElement, parse_xml, nsdecls, qn, Pt, RGBColor, Mm, Inches

    try:
        from docx import Document as _Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as _WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import OxmlElement as _OxmlElement, parse_xml as _parse_xml
        from docx.oxml.ns import nsdecls as _nsdecls, qn as _qn
        from docx.shared import Inches as _Inches, Mm as _Mm, Pt as _Pt, RGBColor as _RGBColor
    except ModuleNotFoundError as exc:
        raise SystemExit(
            "Missing dependency for DOCX export. Install requirements with "
            "`python -m pip install -r markup-hub/requirements.txt`."
        ) from exc

    Document = _Document
    WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
    WD_CELL_VERTICAL_ALIGNMENT = _WD_CELL_VERTICAL_ALIGNMENT
    WD_PARAGRAPH_ALIGNMENT = _WD_PARAGRAPH_ALIGNMENT
    OxmlElement = _OxmlElement
    parse_xml = _parse_xml
    nsdecls = _nsdecls
    qn = _qn
    Pt = _Pt
    RGBColor = _RGBColor
    Mm = _Mm
    Inches = _Inches

    input_path = Path(args.input)
    output_path = Path(args.output)
    template_path = Path(args.template)

    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise SystemExit(
            "Pandoc is required for DOCX export but was not found on PATH. "
            "Install Pandoc from https://pandoc.org/installing.html and rerun."
        )
    if not input_path.exists():
        raise SystemExit(f"Input Markdown not found: {input_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    metadata_text, prepared_markdown = build_pandoc_inputs(
        input_path.read_text(encoding="utf-8")
    )

    temp_input_path: Path | None = None
    temp_metadata_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", encoding="utf-8", suffix=".md", delete=False
        ) as temp_markdown:
            temp_markdown.write(prepared_markdown)
            temp_input_path = Path(temp_markdown.name)

        if metadata_text:
            with tempfile.NamedTemporaryFile(
                mode="w", encoding="utf-8", suffix=".yaml", delete=False
            ) as temp_metadata:
                temp_metadata.write(metadata_text)
                temp_metadata_path = Path(temp_metadata.name)

        command = [
            pandoc,
            str(temp_input_path),
            "--from=markdown-yaml_metadata_block+tex_math_dollars+bracketed_spans",
            "--to=docx",
            "--standalone",
            "--wrap=none",
            "--output",
            str(output_path),
        ]
        if temp_metadata_path is not None:
            command.extend(["--metadata-file", str(temp_metadata_path)])
        if template_path.exists():
            command.extend(["--reference-doc", str(template_path)])

        result = subprocess.run(command, capture_output=True, text=True, check=False)
        if result.returncode != 0:
            message = result.stderr.strip() or result.stdout.strip() or "Pandoc export failed."
            raise SystemExit(message)

        postprocess_docx(output_path)
    finally:
        if temp_input_path is not None:
            temp_input_path.unlink(missing_ok=True)
        if temp_metadata_path is not None:
            temp_metadata_path.unlink(missing_ok=True)

    print(f"Wrote DOCX output to {output_path}")
    if template_path.exists():
        print(f"Applied template styling from {template_path}")
    else:
        print(f"Template not found at {template_path}; used Pandoc output plus local styling defaults.")
    if args.audit_bold_leads:
        print("Bold-lead audit: ignored in the Pandoc-based exporter.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
